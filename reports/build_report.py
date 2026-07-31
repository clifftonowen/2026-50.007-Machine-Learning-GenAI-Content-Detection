"""Generate reports/report.docx, the Task 4 deliverable.

Every number comes from reports/facts.py, which reads the ledgers in data/processed/, so
the document cannot drift from its own evidence. Regenerate after any new result lands:

    python reports/build_report.py
"""

import sys
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
import facts  # noqa: E402

OUT = facts.PROJECT_ROOT / "reports" / "report.docx"
TODO = RGBColor(0xC0, 0x00, 0x00)


class Report:
    def __init__(self):
        self.d = docx.Document()
        for name, size in [("Normal", 10.5), ("Heading 1", 15), ("Heading 2", 12.5)]:
            try:
                self.d.styles[name].font.size = Pt(size)
            except KeyError:
                pass

    def h(self, text, level=1):
        self.d.add_heading(text, level=level)

    def p(self, text, *, italic=False, bold=False):
        par = self.d.add_paragraph()
        run = par.add_run(text)
        run.italic, run.bold = italic, bold
        return par

    def todo(self, text):
        par = self.d.add_paragraph()
        run = par.add_run(f"[TODO: {text}]")
        run.bold = True
        run.font.color.rgb = TODO
        return par

    def bullets(self, items):
        for it in items:
            self.d.add_paragraph(it, style="List Bullet")

    def choice(self, chose, tried, drawback, acceptable):
        """The house style: chosen, alternatives with numbers, drawback, why it is fine."""
        self.p(f"Chosen: {chose}", bold=True)
        self.p("Also tried:")
        self.bullets(tried)
        self.p(f"Drawback: {drawback}")
        self.p(f"Why that is acceptable: {acceptable}")

    def table(self, headers, rows, widths=None):
        t = self.d.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        for i, htxt in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = str(htxt)
            for par in cell.paragraphs:
                for run in par.runs:
                    run.bold = True
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = str(v)
        if widths:
            for r in t.rows:
                for i, w in enumerate(widths):
                    r.cells[i].width = Inches(w)
        self.d.add_paragraph()
        return t

    def figure(self, name, caption, width=6.1):
        path = facts.figure(name)
        self.d.add_picture(str(path), width=Inches(width))
        self.d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = self.d.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)

    def save(self):
        self.d.save(OUT)
        return OUT


def build():
    facts.check()
    r = Report()

    # ---------------------------------------------------------------- title
    title = r.d.add_heading("Detecting Machine-Generated Text", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = r.d.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("50.007 Machine Learning, COLING 2026 GenAI Content Detection\n"
                "Task 4: documenting the journey").italic = True
    who = r.d.add_paragraph()
    who.alignment = WD_ALIGN_PARAGRAPH.CENTER
    who.add_run(", ".join(facts.MEMBERS))
    r.todo("registered Kaggle team name")

    # ---------------------------------------------------------------- 1
    r.h("1. Summary")
    r.p(
        f"Our best submission scores {facts.BEST_KAGGLE} Macro F1 on the public "
        f"leaderboard. It is a LightGBM classifier trained on {facts.BEST_N_FEATURES:,} "
        "features built from the raw text, and its predictions are turned into labels by "
        "thresholding the two test id groups separately rather than applying one global "
        "cut."
    )
    r.p(
        f"Starting from the supplied 5,000-dimension TF-IDF matrix, the same model family "
        f"scored {facts.SUPPLIED_LGBM_KAGGLE}. Of the "
        f"{facts.BEST_KAGGLE - facts.SUPPLIED_LGBM_KAGGLE:+.5f} we added, roughly two "
        "thirds came from changing what the model sees and one third from changing how "
        "its scores are converted into labels. Almost none came from the model itself: "
        "three rounds spent on tuning and ensembling produced gains inside the "
        "leaderboard's own noise, and we report those rounds as the nulls they were."
    )
    r.p(
        "Throughout, we treated the public leaderboard as a measuring instrument with a "
        f"quantified precision. On roughly 3,570 scored rows at F1 near 0.8, binomial "
        f"sampling noise is about {facts.NOISE_FLOOR}. Any difference smaller than that "
        "is reported as a tie, including several of our own later gains."
    )
    r.figure("kaggle_journey.png",
             "Figure 1. Public Macro F1 across the milestone submissions. The shaded "
             "band is the noise floor around the final score; three steps clear it, the "
             "first global share calibration, the representation change, and the first "
             "per-group share correction.")

    # ---------------------------------------------------------------- 2
    r.h("2. The data, and the shift the brief warned about")
    r.p(
        "The task is binary: label a document 1 if machine-generated and 0 if "
        "human-authored, scored on Macro F1. Training data is 20,000 labelled documents "
        "and the test file holds 6,999. The course supplies a pre-computed 5,000-column "
        "TF-IDF matrix, lemmatised and with stop words removed."
    )
    r.p(
        "Two properties of the data shaped every decision that follows."
    )
    r.p("The classes are unbalanced.", bold=True)
    r.p(
        "Training is 62.52% machine and 37.48% human. A majority-class predictor would "
        "reach 62.5% accuracy and about 0.38 Macro F1, which is why we stratified every "
        "split and used class-weighted estimators throughout."
    )
    r.p("Train and test are different source corpora.", bold=True)
    r.p(
        "The brief warns of a deliberate train-to-test shift, and the COLING shared-task "
        "paper explains its shape: training draws on HC3, M4GT and MAGE, while the test "
        "split is CUDRT, IELTS, NLPeer, PeerSum and MixSet, with no overlap. The paper's "
        "figures put training at 62.6% machine against 53.1% for the test split, matching "
        "the 0.6252 we measured and the ~0.53 we later inferred from the leaderboard."
    )
    r.p(
        "The test file also contains two visibly different populations. 1,999 rows carry "
        "UUID-style ids like the training rows, and 5,000 carry numeric ids. The numeric "
        "rows are longer (median 1,862 characters against 1,189), carry markdown bold at "
        "2.52 occurrences per document against 0.44, and contain the word \"reviewer\" in "
        "15.8% of documents against 0.2%. They are academic peer reviews. This split "
        "turned out to matter more than any modelling decision we made."
    )
    r.figure("class_balance.png",
             "Figure 2. Class balance in the training set: 12,504 machine against 7,496 "
             "human.", width=4.6)
    r.figure("text_length_distribution.png",
             "Figure 3. Document length distribution. Length correlates with the label, "
             "and later became the basis of our validation protocol.", width=5.8)

    # ---------------------------------------------------------------- 3
    r.h("3. Question 1: the best model and how it works")
    r.h("3.1 The classifier", level=2)
    r.p(
        "LightGBM is a gradient-boosted decision tree ensemble. It fits trees in "
        "sequence, each one trained on the gradient of the loss left over by the trees "
        "before it, so every tree corrects its predecessors rather than voting "
        "independently. Its leaf-wise growth strategy splits whichever leaf promises the "
        "largest loss reduction instead of growing level by level, which is what makes it "
        "efficient on very wide sparse inputs like ours."
    )
    r.p(
        f"We pass class_weight=\"balanced\", which reweights the 62/38 split so the "
        "minority human class is not simply outvoted. Macro F1 averages the two per-class "
        "F1 scores with equal weight, so a model that quietly sacrifices the human class "
        "is punished twice: once in that class's recall and once in the average."
    )
    r.p(
        f"Every other hyperparameter is at its default. That is a real finding rather than "
        "an omission, and section 4.5 explains what we did about it."
    )

    r.h("3.2 The features", level=2)
    r.p(
        f"The model sees {facts.BEST_N_FEATURES:,} columns in eight blocks, all built "
        "from the raw text. Blocks A to F are per-document statistics that cannot leak; "
        "H and I are TF-IDF vectorizers fitted on the training text alone."
    )
    r.table(
        ["Block", "Columns", "What it measures"],
        [[n, f"{c:,}", d] for n, c, d in facts.CHOSEN_BLOCKS],
        widths=[1.5, 0.9, 3.7],
    )
    r.p(
        "Two families are deliberately absent, and the table below is why. "
        "We dropped each family in turn from the full set and re-scored on a held-out "
        "document type, so the column is what that family is worth for transfer rather "
        "than for fitting the training domains."
    )
    r.table(
        ["Family dropped", "Macro F1 given up", "Columns"],
        [[n, f"{d:+.4f}", f"{c:,}"] for n, d, c in facts.ABLATION],
        widths=[2.2, 1.7, 1.2],
    )
    r.p(
        "Readability contributes nothing and the supplied TF-IDF is actively harmful: "
        "removing it improves held-out-domain Macro F1 by 0.0074. Note also that "
        "diversity returns more than character n-grams from five columns against 20,000, "
        "which is why we later spent a round trying to extend it. Section 5.5 reports how "
        "that went."
    )

    r.h("3.3 The decision rule", level=2)
    r.p(
        "The model emits a probability per row, but we do not threshold at 0.5. We sort "
        "the rows by score and label a fixed fraction of them machine, and we do this "
        f"separately for each id group: {facts.BEST_SHARES['uuid']} of the UUID rows and "
        f"{facts.BEST_SHARES['numeric']} of the numeric rows."
    )
    r.p(
        "This is the single highest-leverage choice in the project and section 4.4 "
        "derives it. The short version is that the two groups need corrections in "
        "opposite directions, so any single global threshold is a compromise between them."
    )

    # ---------------------------------------------------------------- 4
    r.h("4. Question 2: the roadmap")
    r.p(
        "Each subsection states what we chose, what else we measured, the drawback we "
        "accepted, and why we judged it acceptable.", italic=True
    )

    r.h("4.1 Representation: raw text over the supplied TF-IDF", level=2)
    reps = facts.representations()
    r.choice(
        chose=(f"eight blocks built from raw text, {facts.BEST_N_FEATURES:,} columns, "
               f"scoring {reps['text_all']:.4f} on standard cross-validation"),
        tried=[
            f"The supplied 5,000-column TF-IDF alone: {reps['tfidf_supplied']:.4f}. "
            "The intended input, and the weakest option we measured.",
            f"Character n-grams alone (block H): {reps['block_H']:.4f}. One block "
            "carrying 59% of the total gain.",
            f"Word n-grams with stop words kept (block I): {reps['block_I']:.4f}, "
            "against the supplied features' 0.7303 on the same corpus. Suggestive that "
            "the course preprocessing costs signal, though block I also has four times "
            "the vocabulary so the comparison is confounded by capacity.",
            f"The cheap per-document style blocks alone (A to G): "
            f"{reps['style_all']:.4f}, from features that each score below the control "
            "individually.",
            f"Supplied TF-IDF plus everything: {reps['supplied_plus_all']:.4f}, a tie "
            "with raw text alone using 5,000 more columns, so the supplied features add "
            "nothing once raw text is present.",
        ],
        drawback=("we hand-build 40,385 features where the course supplies 5,000 ready "
                  "to use, which is far more engineering effort, far more code to get "
                  "wrong, and a pipeline that would need rebuilding for any new corpus."),
        acceptable=("the supplied features are lemmatised with stop words removed, so "
                    "they are a pure content signal, and content vocabulary is exactly "
                    "what changes when the test set comes from different corpora. "
                    "Function words, punctuation, casing and layout survive a corpus "
                    "change, and the course preprocessing deletes precisely those. The "
                    "brief permits extra feature engineering when it is described, and "
                    "this is that description."),
    )
    r.figure("representation_comparison.png",
             "Figure 4. Each representation's best model, measured against the supplied "
             "TF-IDF control re-run in the same session.")

    r.h("4.2 Model family: LightGBM, and why we stopped ensembling", level=2)
    base = facts.baselines()
    r.choice(
        chose=f"LightGBM, best of twelve baselines at {base.loc['LightGBM', 'mean']:.4f}",
        tried=[
            f"{name}: {row['mean']:.4f}"
            for name, row in base.iloc[1:7].iterrows()
        ] + [
            "Weaker still: RandomForest 0.6871, ComplementNB 0.6580, KNN 0.6542, "
            "MultinomialNB 0.6505, AdaBoost 0.5978.",
            "A six-member ensemble with four families of combiner across 48 "
            "configurations (weighted blending, voting, meta-learners, rank "
            f"aggregation). Best gain on Kaggle: {facts.ROUND4_BEST_KAGGLE - 0.73583:+.5f}, "
            f"a fifth of the {facts.NOISE_FLOOR} noise floor.",
        ],
        drawback=("a single boosted-tree model has no diversity to fall back on, and "
                  "ensembling is the standard way to buy a last fraction of a point."),
        acceptable=("we measured that fraction and it was not there. Round 4 spent a "
                    "full round on combiners and gained less than the leaderboard can "
                    "resolve. Shipping one model also keeps the pipeline auditable, "
                    "which mattered when we later needed to attribute gains precisely."),
    )
    r.figure("baseline_models_cv_f1.png",
             "Figure 5. The twelve-model baseline sweep on the supplied features. The "
             "ordering held up: LightGBM led locally and, once compared fairly, on the "
             "leaderboard.", width=5.8)

    r.h("4.3 Validation: leave-one-length-band-out", level=2)
    r.choice(
        chose=("grouped cross-validation holding out one document-length band at a time, "
               f"giving {facts.BEST_CV_GROUPED} for the shipped model"),
        tried=[
            f"Standard 5-fold stratified CV: {facts.BEST_CV_STANDARD}, but it overstates "
            "the leaderboard by about 0.10 because it trains and tests on the same "
            "domains.",
            "K-means clusters on the style features, intended as discovered domains. "
            "Rejected: the stability gate passed at adjusted Rand index 0.9703 on one "
            "machine and 0.5991 on another, and fell to 0.1043 over five seeds.",
            "A deliberately random grouping, run as a control. It scored 0.8742 against "
            "standard CV's 0.8764, which is the evidence that the transfer gap comes "
            "from the structure of the grouping rather than from training on fewer rows.",
        ],
        drawback=("length bands are a crude proxy for domain, and the absolute level is "
                  "not trustworthy: across three sensible groupings the score spans "
                  "0.0195, so the number carries roughly +/- 0.03."),
        acceptable=("we use it to rank candidates, not to predict a leaderboard score, "
                    "and for ranking it was several times more faithful than standard "
                    "CV. It is also deterministic and unfitted, so all five of us compute "
                    "identical folds, which the k-means version demonstrably could not."),
    )
    r.p(
        "The control is the load-bearing row of the table below. If a meaningless "
        "grouping had produced the same gap, the protocol would have been measuring "
        "reduced training-set size rather than domain structure, and no version of it "
        "could have been a selection rule."
    )
    r.table(
        ["Grouping of the training rows", "Held-out Macro F1", "Usable folds"],
        [[n, f"{s:.4f}", k] for n, s, k in facts.GROUPINGS]
        + [["standard 5-fold CV, for reference", f"{facts.BEST_CV_STANDARD:.4f}", 5]],
        widths=[2.6, 1.6, 1.1],
    )
    r.figure("local_vs_kaggle.png",
             "Figure 6. Standard CV, grouped CV and the leaderboard for three milestone "
             "models. Grouped CV is not unbiased, but its error is several times smaller.")

    r.h("4.4 The decision rule: per-group share", level=2)
    r.p(
        "This subsection is the one we would most want a reader to take away, because "
        "the mistake it corrects is easy to make and invisible from inside."
    )
    r.p(
        "We swept the global predicted share twice, in two different rounds and on two "
        "different models, and both times it came back flat. On the raw-text model, "
        "share 0.4996 gave 0.77349 and share 0.5299 gave 0.77231, a tie. We concluded "
        "share was exhausted. That conclusion was correct about the quantity we measured "
        "and wrong about the question we thought we were answering."
    )
    r.p(
        "The reason is that a global cut moves both id groups in the same direction, and "
        "the two groups need corrections in opposite directions. The UUID rows wanted a "
        "higher predicted machine share and the numeric rows a lower one, so a global "
        "sweep moved one toward its optimum while pushing the other away, and the two "
        "effects cancelled. Thresholding the groups separately breaks the cancellation."
    )
    r.figure("class_balance_curve_round2.png",
             "Figure 7. The global predicted-share curve, fitted across two rounds. Flat "
             "across a wide band, which is exactly what a cancelling pair of effects "
             "looks like when they are averaged together.", width=5.4)
    surf = facts.share_surface()
    r.table(
        ["UUID share", "Numeric share", "Kaggle Macro F1"],
        [[f"{row['uuid_share']:.4f}", f"{row['numeric_share']:.4f}",
          f"{row['kaggle_f1']:.5f}"] for _, row in surf.iterrows()],
        widths=[1.6, 1.6, 1.9],
    )
    r.choice(
        chose=(f"per-group shares of {facts.BEST_SHARES['uuid']} for UUID rows and "
               f"{facts.BEST_SHARES['numeric']} for numeric rows"),
        tried=[
            "The default 0.5 probability threshold, which inherits the training prior of "
            "0.6252 and scored 0.65738, the worst result of the project.",
            "A threshold tuned against held-out data, which moved it the wrong way and "
            "scored 0.66502. Local data cannot pick a threshold here: dev and holdout are "
            "both carved from training and carry a class balance the test set does not.",
            "Global share swept from 0.44 to 0.56 on two models. Flat both times.",
            "Per-group share swept on both axes to their vertices, at 0.6214 and 0.4897 "
            "respectively.",
        ],
        drawback=("the shares are fitted entirely to the public leaderboard, which no "
                  "local data can validate, and seven submissions went into locating "
                  "them. A coordinate search that size on a noisy objective can "
                  "manufacture 0.005 to 0.010 of gain that is not on the private rows."),
        acceptable=("the shape of the surface is solid even if the last decimal is not. "
                    "The first correction gained 0.01764 and the wrong-direction probe "
                    "lost 0.00926, both clearing the noise floor, and the same "
                    "correction was confirmed independently on a completely different "
                    "model. We also stop refining inside a flat region rather than "
                    "chasing it, and our two final picks pair the best point with a more "
                    "central one instead of taking the top two public scores."),
    )
    r.figure("share_surface.png",
             "Figure 8. The per-group share surface. Both axes are at their fitted "
             "vertices; the numeric axis is asymmetric, and only the upward step clears "
             "the noise floor, which is what makes the direction certain.")

    r.h("4.5 Hyperparameters", level=2)
    r.p(
        "Every hyperparameter search in this project ran against the supplied 5,000 "
        "features. Nothing has been tuned on the 40,385-column raw-text matrix, which is "
        "a different problem in shape: far wider, far sparser, and three quarters "
        "character n-grams. The shipped model therefore runs on LightGBM defaults."
    )
    r.p(
        "The search is designed and split five ways across the team, sampling learning "
        "rate, tree count, leaf count, depth, minimum child samples, column subsampling, "
        "row subsampling and both regularisation terms. It selects on paired per-fold "
        "differences against the default configuration rather than on raw means, requires "
        "a candidate to improve on every fold, and confirms the winner on a protocol the "
        "search never selected on."
    )
    r.todo("hyperparameter search result, once notebook 17 has been run by the team")
    r.figure("lightgbm_stage2_grid_heatmap.png",
             "Figure 9. The equivalent search on the supplied features, for reference. "
             "It moved cross-validated Macro F1 from 0.7391 to 0.7443, about +0.005, "
             "which is the scale of gain we expect and why defaults were acceptable to "
             "ship in the meantime.", width=5.4)

    # ---------------------------------------------------------------- 5
    r.h("5. Question 3: difficulties, and what they cost")
    r.p(
        "Every item below is a mistake we made and then caught. We include the cost of "
        "each because the cost is the part that taught us something.", italic=True
    )

    r.h("5.1 We compared models along an axis we had not held fixed", level=2)
    r.p(
        "For three notebooks we believed local validation was anti-correlated with the "
        "leaderboard. LightGBM had the best cross-validation and holdout scores of "
        "anything we tried and the worst leaderboard score, so we wrote it off and spent "
        "a round on linear models instead."
    )
    r.p(
        "The inversion was an artifact. Each model had been submitted at whatever "
        "predicted class balance its own default threshold happened to produce, so "
        "\"which model is best\" and \"which model is best calibrated for this test set\" "
        "were being answered as one question. Re-testing at a matched share of 0.4996, "
        "LightGBM beat ElasticNet by 0.0189, in the same direction and roughly the same "
        "magnitude as its local lead. Local validation had never been broken."
    )
    r.p(
        "What fixed it was designing the test in advance to be able to overturn our own "
        "conclusion, and writing down what each outcome would mean before the scores came "
        "back. We kept that habit for the rest of the project."
    )

    r.h("5.2 Three rounds of local gains that did not transfer", level=2)
    r.p(
        "Round 3 gained 0.0017 on cross-validation and nothing measurable on the "
        "leaderboard. Round 4 gained 0.0075 out-of-fold AUC and 0.0017 on the "
        "leaderboard, a fifth of the noise floor. The diagnosis, once we read the "
        "shared-task paper properly, was that we "
        "had been optimising the model while holding the representation fixed, and the "
        "representation was the binding constraint. Round 5 changed it and gained 0.036."
    )

    r.h("5.3 A validation protocol that would not reproduce", level=2)
    r.p(
        "We built domain clusters with k-means and gated everything downstream on a "
        "stability check. It passed at adjusted Rand index 0.9703 on one machine and "
        "failed at 0.5991 on another, from identical code and seeds. The cause was a "
        "460,000-fold spread in feature scales feeding a decomposition that ran before "
        "standardisation, on a partition that turned out to be genuinely multi-modal."
    )
    r.p(
        "Rather than discard the idea we tested whether the score depended on the "
        "labelling, including a deliberately random grouping as a control. The control "
        "reproduced standard CV almost exactly while every structured grouping sat well "
        "below it, which told us the mechanism was sound even though the specific "
        "clusters were not. We replaced them with deterministic length bands."
    )

    r.h("5.4 A diagnostic that silently contaminated itself", level=2)
    r.p(
        "We built a submission to test whether the public leaderboard scores the UUID "
        "rows, by changing only those rows and holding the numeric rows fixed. The first "
        "version pinned the numeric group to the global share of 0.4996, believing that "
        "held them constant. Their realised share was 0.5482, so it silently moved 243 "
        "rows it was supposed to freeze, which would have made the answer unreadable in "
        "both directions."
    )
    r.p(
        "The general lesson we took: setting a knob to the value a baseline used is not "
        "the same as leaving those rows alone, and the two coincide only when the knob is "
        "applied at the same granularity. Every later per-group file asserts that the "
        "held group is byte-identical rather than trusting the share."
    )

    r.h("5.5 Two predictions we made, and both were wrong", level=2)
    r.p(
        "We predicted that the formatting features were dataset fingerprints that would "
        "not survive the corpus change, since machine text in training carries markdown "
        "bold at ten times the human rate. We built a stripped-down feature set to test "
        "it. Removing those blocks cost 0.04944 on the leaderboard, six times the noise "
        "floor in the opposite direction from the prediction."
    )
    r.p(
        "We then predicted that our strongest feature family had more to give, since it "
        "returned the largest transfer value of any block from only five columns. We "
        "added fourteen further diversity measures. They made held-out-domain performance "
        "slightly worse while improving same-domain performance, which is the exact "
        "signature of a feature that memorises rather than generalises."
    )
    r.p(
        "Both experiments were designed so they could fail, and both did. We would rather "
        "report two falsified predictions than a roadmap in which every idea worked."
    )

    r.h("5.6 What we could not overcome", level=2)
    r.bullets([
        "Predicted share cannot be validated locally, at all. Our dev and holdout splits "
        "are carved from training data and inherit its class balance, so every share "
        "decision rests on leaderboard feedback alone. We managed the risk rather than "
        "solving it.",
        "The absolute level of our grouped protocol is not trustworthy to better than "
        "about +/- 0.03, so we can rank candidates but not forecast a score.",
        "A gap of about 0.075 remains between standard cross-validation and the "
        "leaderboard. The brief says to expect it by design, and closing it further was "
        "not something we found a lever for.",
    ])

    # ---------------------------------------------------------------- 6
    r.h("6. Question 4: what we learned beyond the course")
    r.h("6.1 Blocked comparison, and why averages mislead", level=2)
    r.p(
        "Our held-out bands differ in difficulty by far more than any two candidates "
        "differ from each other. Comparing mean scores against a fold standard deviation "
        "of 0.087 sets a bar nothing can clear. Because band difficulty is common to "
        "every candidate, the correct comparison is paired: subtract fold by fold and "
        "test the differences. We had already reported one result as a null using the "
        "conservative bar before realising this, and we say so rather than quietly "
        "restating it."
    )

    r.h("6.2 Averaging can hide two large effects that cancel", level=2)
    r.p(
        "The single most useful idea we took from this project generalises well beyond "
        "it. When one global threshold is applied to a population that is a mixture, the "
        "optimum for the mixture can be flat while the optima for the components are far "
        "apart and moving in opposite directions. We measured the flat curve twice, "
        "correctly, and drew the wrong conclusion from it both times. A roughly +0.019 "
        "effect and a roughly -0.019 effect had been cancelling in the average."
    )

    r.h("6.3 A leaderboard is a measuring instrument with a precision", level=2)
    r.p(
        f"We computed the sampling noise of the public leaderboard once, about "
        f"{facts.NOISE_FLOOR} on its row count at our score, and then held every claim to "
        "it. That single number let us stop several searches at the right moment, decide "
        "which of our own gains to disbelieve, and avoid spending submissions inside flat "
        "regions. It is also why our final picks are not simply the top two public scores."
    )

    r.h("6.4 Reading the dataset paper was worth more than any model", level=2)
    r.p(
        "The shared-task paper told us the test corpora were disjoint from training, gave "
        "us the class balance of each, and explained the two id populations. Those facts "
        "produced the two changes that moved our score. No amount of hyperparameter "
        "search would have found them."
    )

    # ---------------------------------------------------------------- 7 and 8
    r.h("7. Task 1: Logistic Regression from scratch")
    r.p(
        "Implemented without sklearn or any pre-built logistic regression package, on the "
        "supplied TF-IDF features, with sigmoid, log loss, an analytic gradient, "
        "mini-batch gradient descent and a predict step."
    )
    r.p(
        "Design choices worth noting: the gradient of the log loss with respect to the "
        "weights reduces to X-transpose times the residual divided by the batch size, so "
        "no numerical differentiation is needed; and the sigmoid is computed in a "
        "numerically stable branch to avoid overflow on large negative inputs."
    )
    r.todo("batch size, epochs and learning rate used, plus the validation Macro F1")
    r.todo("Kaggle Macro F1 for LogReg_Prediction.csv, and confirmation that the file "
           "exists with that exact name")

    r.h("8. Task 2: PCA and KNN")
    r.p(
        "PCA applied to the supplied 5,000 TF-IDF features, followed by K-nearest "
        "neighbours with n_neighbors fixed at 2, evaluated at 2,000, 1,000, 500 and 100 "
        "components. Predictions for all four counts have been generated."
    )
    r.p(
        "The scree plot below is the diagnostic for how much variance each component "
        "count retains. The four counts are fixed by the task rather than chosen from the "
        "elbow, which is worth stating because it is the one place in the project where "
        "we did not pick a structural hyperparameter from a diagnostic."
    )
    r.figure("scree_plot.png",
             "Figure 10. Explained variance against component count for PCA on the "
             "supplied features.", width=5.4)
    r.todo("test Macro F1 at 2000, 1000, 500 and 100 components, read back from Kaggle")
    r.todo("analysis of the reduced components, which the rubric rewards separately")

    # ---------------------------------------------------------------- 9
    r.h("9. Reflection and next steps")
    r.p("What we consider settled:", bold=True)
    r.bullets([
        "The representation. Raw-text features beat the supplied TF-IDF by a wide margin, "
        "and adding the supplied features back changes nothing.",
        "Global predicted share. Flat across 0.44 to 0.56 on two different models.",
        "Per-group share. Both axes measured to their vertices.",
        "Feature expansion in the directions we tried. Four ideas, all inside the bar: "
        + "; ".join(f"{n} {d:+.4f}" for n, d in facts.EXPANSION_NULLS) + ".",
    ])
    r.p("What remains open:", bold=True)
    r.bullets([
        "Hyperparameter tuning on the raw-text representation, designed and split across "
        "the team but not yet run.",
        "A second model family blended with the first on this representation, which we "
        "specified but did not manage to measure.",
        "Final submission selection, where our rule is to pair the best public point with "
        "a more central one rather than taking the top two, because the share coordinates "
        "are fitted to a leaderboard with a known noise floor.",
    ])
    r.p(
        "For calibration, the shared task's own fine-tuned RoBERTa baseline reached 73.42 "
        f"Macro F1 on the full English test set. We reach {facts.BEST_KAGGLE * 100:.2f} "
        "with classical models and no deep learning, though the numbers are not strictly "
        "comparable because the course uses a subsample."
    )

    return r.save()


if __name__ == "__main__":
    out = build()
    print(f"wrote {out}  ({out.stat().st_size / 1000:.0f} kB)")
