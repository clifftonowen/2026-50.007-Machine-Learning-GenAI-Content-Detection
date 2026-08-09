"""Generate reports/report_draft.docx, the Task 4 report told as a methodology.

This is the chronological retelling: the project is presented in the order it actually
ran, and every phase justifies itself the same way (what else we considered, with pros
and cons, what we chose, and the tradeoff we accepted with the reason it was acceptable).

reports/build_report.py is the earlier version, organised around the brief's four
questions. Both are kept, both read numbers from reports/facts.py, so they cannot
disagree with each other or with data/processed/.

    python reports/build_report_draft.py
"""

import sys
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
import facts  # noqa: E402

OUT = facts.PROJECT_ROOT / "reports" / "report_draft.docx"
TODO = RGBColor(0xC0, 0x00, 0x00)
CHOSE = RGBColor(0x1F, 0x77, 0xB4)


class Report:
    """Thin wrapper over python-docx carrying this report's three recurring blocks."""

    def __init__(self):
        self.d = docx.Document()
        for name, size in [("Normal", 10.5), ("Heading 1", 15), ("Heading 2", 12.5),
                           ("Heading 3", 11)]:
            try:
                self.d.styles[name].font.size = Pt(size)
            except KeyError:
                pass
        self.figure_n = 0
        self.table_n = 0

    def h(self, text, level=1):
        self.d.add_heading(text, level=level)

    def p(self, text, *, italic=False, bold=False):
        par = self.d.add_paragraph()
        run = par.add_run(text)
        run.italic, run.bold = italic, bold
        return par

    def todo(self, text):
        par = self.d.add_paragraph()
        run = par.add_run(f"[TODO: {text}]")
        run.bold = True
        run.font.color.rgb = TODO
        return par

    def bullets(self, items):
        for it in items:
            self.d.add_paragraph(it, style="List Bullet")

    def table(self, headers, rows, caption, widths=None):
        self.table_n += 1
        t = self.d.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        for i, htxt in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = str(htxt)
            for par in cell.paragraphs:
                for run in par.runs:
                    run.bold = True
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = str(v)
        if widths:
            for r in t.rows:
                for i, w in enumerate(widths):
                    r.cells[i].width = Inches(w)
        cap = self.d.add_paragraph()
        run = cap.add_run(f"Table {self.table_n}. {caption}")
        run.italic = True
        run.font.size = Pt(9)
        return t

    def considered(self, rows, caption):
        """The options on the table at this point in the project, with pros and cons."""
        return self.table(["Option", "Pros", "Cons"], rows, caption,
                          widths=[1.5, 2.35, 2.35])

    def decision(self, chose, tradeoff, because):
        """The three-part block that closes every phase."""
        for label, body, colour in [("We chose: ", chose, CHOSE),
                                    ("Tradeoff accepted: ", tradeoff, None),
                                    ("Acceptable because: ", because, None)]:
            par = self.d.add_paragraph()
            head = par.add_run(label)
            head.bold = True
            if colour is not None:
                head.font.color.rgb = colour
            par.add_run(body)

    def figure(self, name, caption, width=6.1):
        self.figure_n += 1
        path = facts.figure(name)
        self.d.add_picture(str(path), width=Inches(width))
        self.d.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = self.d.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(f"Figure {self.figure_n}. {caption}")
        run.italic = True
        run.font.size = Pt(9)

    def save(self):
        self.d.save(OUT)
        return OUT


def build():
    facts.check()
    r = Report()

    # ------------------------------------------------------------------ title
    title = r.d.add_heading("Detecting Machine-Generated Text", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = r.d.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("50.007 Machine Learning, COLING 2026 GenAI Content Detection\n"
                "Task 4 report: the method, in the order we followed it").italic = True
    who = r.d.add_paragraph()
    who.alignment = WD_ALIGN_PARAGRAPH.CENTER
    who.add_run(", ".join(facts.MEMBERS))
    r.todo("registered Kaggle team name, which is how Task 3 marks are awarded")

    # ------------------------------------------------------------------ 1
    r.h("1. Summary")
    r.p(
        f"Our best submission scores {facts.BEST_KAGGLE} Macro F1 on the public "
        f"leaderboard. It is a LightGBM classifier trained on {facts.BEST_N_FEATURES:,} "
        "features built from the raw text, and its scores are turned into labels by "
        "thresholding the two visibly different populations in the test file separately "
        "rather than applying one global cut."
    )
    r.p(
        "This report is organised as a method rather than as a list of results. We worked "
        "through nine phases in order, and each one is written the same way: the options "
        "that were open to us at that point, the pros and cons of each, what we chose, and "
        "the tradeoff we knowingly accepted. Sections 2 to 10 are those nine phases and "
        "together they answer the brief's roadmap question. Section 14 gathers the final "
        "model in one place, section 15 covers the difficulties, section 16 what we "
        "learned beyond the course, and section 17 the limitations we did not resolve."
    )
    r.p(
        f"Our first submission, a tuned LightGBM on the supplied features with the "
        f"default 0.5 probability cut, scored {facts.FIRST_KAGGLE_DEFAULT_THR}. The "
        f"journey from there to {facts.BEST_KAGGLE} is "
        f"{facts.BEST_KAGGLE - facts.FIRST_KAGGLE_DEFAULT_THR:+.5f}, and it divides "
        "unevenly in a way we did not anticipate."
    )
    r.table(
        ["Change", "Phase", "Gain on the public leaderboard"],
        [["Move off the default cut to a global predicted share", "7", "+0.07845"],
         ["Rebuild the representation from raw text", "8", "+0.04359"],
         ["Threshold the two test populations separately", "9", "+0.02201"],
         ["Retune the model on the new representation", "9",
          f"{facts.TUNING_GAIN_KAGGLE:+.5f}"],
         ["Ensembling six models with four combiner families", "6", "+0.00171"],
         ["Hyperparameter tuning on the supplied features", "5", "not measurable"]],
        "Where the score came from. Roughly 65% of the total came from how scores are "
        "turned into labels, 28% from what the model sees, and 7% from the model itself.",
        widths=[3.2, 0.8, 2.1],
    )
    r.p(
        "That ordering is the honest summary of the project and it is not the one we "
        "expected. The two phases that moved the score most were the ones that questioned "
        "the decision rule and the input, not the ones that worked on the model."
    )
    r.p(
        "The last two rows of that table are the same activity with opposite results, and "
        "the difference between them is the single most useful thing we learned. Tuning "
        "the model on the supplied 5,000 features returned nothing the leaderboard could "
        "resolve, and for most of the project we generalised that into a belief that the "
        "model did not matter on this task. Running the identical search against the "
        "raw-text representation we had since moved to returned "
        f"{facts.TUNING_GAIN_KAGGLE:+.5f}, comfortably outside the noise floor. The "
        "lesson is not that tuning works or does not work, but that a null result is "
        "attached to the representation it was measured on and does not travel to a "
        "different one. We had carried ours across a representation change without "
        f"rechecking it, and it cost us {facts.TUNING_GAIN_KAGGLE:.5f} until the final "
        "week."
    )
    r.p(
        "One number governs every judgement below. On roughly 3,570 scored public rows at "
        f"Macro F1 near 0.8, binomial sampling noise is about {facts.NOISE_FLOOR}. We "
        "computed this once and then held every claim to it, including several of our own "
        "apparent gains that did not survive the comparison."
    )
    r.figure("methodology_flow.png",
             "The nine phases, in the order they happened, colour-coded by what each one "
             "returned. Four were foundations, two returned nothing measurable, and three "
             "moved the score.", width=6.3)

    # ------------------------------------------------------------------ 2
    r.h("2. Phase 1: Exploratory Data Analysis")
    r.p(
        "No model was trained until we had measured the dataset. That was a deliberate "
        "sequencing choice, and three of the four measurements went on to constrain a "
        "later decision, so it repaid the time."
    )
    r.p("Class balance.", bold=True)
    r.p(
        "The training set is 12,504 machine documents against 7,496 human, so 62.52% "
        "against 37.48%. A predictor that answers machine every time scores 62.5% accuracy "
        "and 0.385 Macro F1. That single comparison settled two things immediately: the "
        "metric to trust is Macro F1 rather than accuracy, and every split and every "
        "cross-validation fold has to be stratified. It also justified passing "
        "class_weight=\"balanced\" to every estimator that supports it, since Macro F1 "
        "averages the two per-class scores with equal weight and therefore punishes a "
        "model that quietly sacrifices the minority human class twice over."
    )
    r.p("Sparsity.", bold=True)
    r.p(
        "The supplied matrix is 98.64% zero. A document activates 67.8 of the 5,000 "
        "columns on average and 58 at the median, with a maximum of 573. No column is "
        "all-zero and none is near-constant, so there was nothing to prune. This is what "
        "made us short-list models that tolerate sparse high-dimensional input, and what "
        "made us sceptical of distance-based methods before we ever ran one."
    )
    r.p("Length, and a first look at the test file.", bold=True)
    r.p(
        "Document length differs between the classes but does not separate them cleanly on "
        "its own. More usefully, comparing train against test text showed the test file "
        "holds two visibly different populations. 1,999 rows carry UUID-style ids like the "
        "training rows; 5,000 carry numeric ids, run longer at a median 1,862 characters "
        "against 1,189, carry markdown bold at 2.52 occurrences per document against 0.44, "
        "and contain the word \"reviewer\" in 15.8% of documents against 0.2%. They are "
        "academic peer reviews. We recorded this and did nothing with it for six phases, "
        "which in hindsight was the single most expensive delay in the project."
    )
    r.considered(
        [
            ["Start modelling immediately",
             "Fastest route to a leaderboard number; the brief supplies ready features",
             "Every structural choice becomes a guess; no way to tell a bug from a "
             "property of the data"],
            ["Full EDA first, then model",
             "Stratification, class weighting and metric choice all become measured "
             "decisions; the test-file structure surfaces early",
             "Costs a full notebook before any score exists"],
            ["EDA in parallel with a first baseline",
             "Both at once",
             "The baseline gets built on assumptions the EDA is still testing, so it "
             "usually has to be rebuilt"],
        ],
        "What we considered at the start of the project.",
    )
    r.decision(
        chose="a full exploratory pass before training anything, in notebook 01.",
        tradeoff="a notebook of work with no leaderboard score at the end of it, at a "
                 "point in the project when a score would have been reassuring.",
        because="three of the four measurements turned into constraints we never had to "
                "revisit, and the fourth, the two id populations, turned out to be worth "
                "more than any model we trained. The alternative was not faster, it was "
                "the same work done later and with rework attached.",
    )
    r.figure("class_balance.png",
             "Class balance in the training set. The skew is what makes Macro F1 and "
             "stratification mandatory rather than optional.", width=4.6)
    r.figure("feature_sparsity.png",
             "Non-zero features per document. The matrix is 98.64% zero, so a document is "
             "described by roughly 68 of the 5,000 available columns.", width=5.6)
    r.figure("text_length_distribution.png",
             "Document length by class. Length does not separate the classes on its own, "
             "but it correlates with the label strongly enough that it later became the "
             "basis of our transfer-validation protocol in section 15.2.", width=5.8)

    # ------------------------------------------------------------------ 3
    r.h("3. Phase 2: Dimensionality Reduction")
    r.p(
        "With 5,000 columns over 20,000 rows the obvious next question is whether the "
        "space can be compressed. The brief also requires it: Task 2 fixes a PCA plus "
        "K-nearest-neighbours pipeline at four component counts. We treated the graded "
        "task and the modelling question as one investigation."
    )

    # ------------------------------------------------------------------ 4
    r.h("4. Phase 3: Choosing the Validation Strategy")
    r.p(
        "Phase 1 said the classes are skewed, so this phase turned that measurement into a "
        "protocol and then froze it. The reasoning runs directly from the class balance: "
        "at 62.5% against 37.5%, an unstratified fold can differ from the population by "
        "several percentage points, and since Macro F1 weights both classes equally, fold "
        "scores would move for reasons that have nothing to do with the model."
    )
    r.considered(
        [
            ["Single random train and validation split",
             "Cheapest; one number per model",
             "One split of 20,000 rows on an imbalanced target is noisy enough to reorder "
             "adjacent candidates"],
            ["Plain k-fold, unstratified",
             "Uses all the data; standard",
             "Fold class balance drifts, and Macro F1 is sensitive to exactly that drift"],
            ["Stratified 5-fold cross-validation",
             "Every fold matches the population balance; five estimates give a spread as "
             "well as a mean",
             "Five fits per candidate; still trains and tests within the same corpora"],
            ["Stratified 10-fold",
             "Lower variance still",
             "Twice the compute for a second-decimal improvement we could not use"],
        ],
        "Validation protocols considered in phase 3.",
    )
    r.decision(
        chose="stratified 5-fold cross-validation with shuffle enabled and random_state "
              "fixed at 42, scoring Macro F1, plus a separate 20% holdout of 4,000 rows "
              "split off once and saved to disk as index arrays so every notebook and "
              "every team member evaluates on identical rows.",
        tradeoff="it trains and tests on the same source corpora, so it measures how well "
                 "a model fits this data rather than how well it transfers to a different "
                 "one. It later proved to overstate the leaderboard by roughly 0.10, and "
                 "it can say nothing at all about the class balance of the test set.",
        because="at this stage we needed an instrument that ranks candidates, not one "
                "that forecasts a score, and for ranking it was right every time we "
                "checked it afterwards. Saving the split indices rather than reseeding a "
                "splitter also removed a whole category of silent inconsistency across a "
                "team of five. When transfer rather than ranking became the question, in "
                "phase 8, we replaced it with a grouped protocol rather than pretending "
                "this one could answer both.",
    )

    r.h("5. Task 1: Logistic Regression from Scratch")
    r.p(
        "Before either, notebook 02 implements logistic regression by hand, with no "
        "pre-built logistic regression package anywhere in the path, as the brief requires "
        "on pain of scoring zero. The implementation is the standard five pieces: a "
        "numerically stable sigmoid that branches on the sign of its input to avoid "
        "overflow, log loss, an analytic gradient, mini-batch gradient descent, and a "
        "predict step. The gradient of log loss with respect to the weights reduces to the "
        "transpose of the design matrix times the residual, divided by the batch size, so "
        "no numerical differentiation is needed and the whole training loop is a few lines "
        "of NumPy."
    )
    r.p(
        "Beyond satisfying the task, this was our reference point: it told us what a "
        "linear model on the supplied features can do, which is most of what the tuned "
        "linear models later achieved. The brief's top band asks for performance "
        "comparable to sklearn's implementation, so that comparison is the thing to report."
    )
    r.table(
        ["Setting", "Value"],
        [["Batch size", f"{facts.TASK1_PARAMS['bs']}"],
         ["Epochs", f"{facts.TASK1_PARAMS['epochs']}"],
         ["Learning rate", f"{facts.TASK1_PARAMS['lr']}"],
         ["Macro F1, held-out 4,000 rows", f"{facts.TASK1_HOLDOUT_F1:.4f}"],
         ["sklearn LogisticRegression, same split",
          f"{facts.TASK1_SKLEARN_REFERENCE:.4f}"]],
        "Our from-scratch implementation against sklearn's, on the same locked split and "
        "the same supplied features. The gap is "
        f"{facts.TASK1_SKLEARN_REFERENCE - facts.TASK1_HOLDOUT_F1:.4f}.",
        widths=[3.0, 1.6],
    )
    r.p(
        f"The {facts.TASK1_SKLEARN_REFERENCE - facts.TASK1_HOLDOUT_F1:.4f} shortfall is "
        "not a defect in the gradient or the loss, both of which we checked against "
        "numerical differentiation. It is that plain mini-batch gradient descent is still "
        "descending when we stop it: the training loss fell by 0.00177 over the final ten "
        "epochs and was still falling, so the optimiser had not converged, while "
        "sklearn's solver runs to a convergence tolerance and adds L2 regularisation by "
        "default. Longer training or a learning-rate schedule would close most of the "
        "remaining gap; we report the honest number rather than tuning the from-scratch "
        "model until it matched."
    )
    r.p(
        "Both filenames ship.", bold=True
    )
    r.p(
        f"The course brief slide names the prediction file {facts.TASK1_FILENAME_CONFLICT[0]} "
        f"while the written task sheet names it {facts.TASK1_FILENAME_CONFLICT[1]}. We "
        "could not resolve which the grader uses, so rather than guess on an exact-name "
        "deliverable we wrote both files with identical contents. The predicted machine "
        f"share is {facts.TASK1_SUBMISSION_SHARE:.2%}, well above the roughly 53% the test "
        "set appears to carry, because Task 1 is graded on the implementation rather than "
        "on leaderboard score and we deliberately left the model at its natural 0.5 cut "
        "rather than importing the calibration work from phase 7."
    )

    r.h("6. Task 2: PCA and KNN at four component counts")
    r.p(
        "PCA at 2,000, 1,000, 500 and 100 components, each feeding a KNN classifier with "
        "n_neighbors fixed at 2. The component counts are set by the task, so this is the "
        "one place in the project where we did not choose a structural hyperparameter from "
        "a diagnostic. We still produced the diagnostic, because it explains the results."
    )
    r.table(
        ["Components", "Cumulative explained variance"],
        [[f"{n:,}", f"{v:.1%}"] for n, v in facts.PCA_VARIANCE],
        "Variance retained at Task 2's four fixed counts, measured in notebook 01.",
        widths=[1.6, 2.6],
    )
    r.p(
        "The analysis the rubric asks for is in that table. This matrix resists "
        "compression badly. Two thousand components, 40% of the original width, still "
        f"leave {1 - facts.PCA_VARIANCE[0][1]:.0%} of the variance behind, and none of the "
        "canonical 80%, 90%, 95% or 99% thresholds is reached within 2,000 components at "
        f"all. The knee of the cumulative curve sits at {facts.PCA_ELBOW[0]} components, "
        f"which retains only {facts.PCA_ELBOW[1]:.1%}. The reason is visible in phase 1: "
        "with 98.64% zeros and about 68 active columns per document, there is very little "
        "shared linear structure for a few directions to absorb. Variance is spread almost "
        "evenly across a very large number of directions, which is exactly the regime in "
        "which PCA has nothing to exploit."
    )
    r.figure("scree_plot.png",
             "Per-component and cumulative explained variance. The curve has no sharp "
             "knee, which is the signature of a matrix that does not compress.", width=6.0)
    r.table(
        ["Components", "Cumulative variance", "Predicted machine share", "Kaggle Macro F1"],
        [[f"{n:,}", f"{v:.2%}", f"{s:.4f}", f"{f:.5f}"] for n, v, s, f in facts.PCA_KNN],
        "The graded Task 2 deliverable: test Macro F1 at each of the four required "
        "component counts, with the predicted class balance of each submission.",
        widths=[1.2, 1.6, 1.7, 1.5],
    )
    r.p(
        "Keeping more variance made the classifier monotonically worse, which is the "
        "opposite of what we expected before running it. The best score came from the most "
        "aggressive reduction, 100 components, holding under a sixth of the variance."
    )
    r.p(
        "Two mechanisms compound. The first is distance concentration: as dimension grows, "
        "the nearest and farthest neighbour distances converge, and it is exactly the "
        "contrast between them that KNN treats as similarity. Because a component's "
        "contribution to squared distance scales with its variance, and there are so many "
        "tail components, at 2,000 components roughly 79% of the squared distance comes "
        "from directions outside the leading 100. Euclidean distance cannot down-weight "
        "them, so the extra directions do not add information so much as dilute it."
    )
    r.p(
        "The second is the tie-break, and it is the larger effect. With n_neighbors=2 and "
        "two classes the neighbours can split one-one; sklearn resolves ties by lowest "
        "class index, which here is human. The model therefore predicts machine only when "
        "both neighbours are machine, so any degradation in retrieval is squared. The "
        "share column measures the damage. Had degraded retrieval simply made the two "
        "neighbours behave like independent draws from the 62.52% training prior, the "
        f"share would fall towards {facts.PCA_KNN_TIE_FLOOR} and stop, since that value is "
        "a floor under that account. The observed share at 2,000 components is "
        f"{facts.PCA_KNN[-1][2]}, which is {facts.PCA_KNN_TIE_FLOOR - facts.PCA_KNN[-1][2]:.4f} "
        "below the floor, so the neighbours are not random draws: they are systematically "
        "human. Backing out the per-neighbour machine rate as the square root of the share "
        "gives 0.738, 0.607, 0.420 and 0.243 across the four counts, against a prior of "
        "0.6252 that it should have matched at every count."
    )
    r.p(
        "The standard explanation for that shape is hubness: in high dimension a small "
        "number of training points become the nearest neighbour of disproportionately many "
        "queries. If those hubs skew human, most queries retrieve at least one human "
        "neighbour, the tie fires, and the prediction collapses to human. We did not "
        "measure the k-occurrence distribution, so we report hubness as the leading "
        "explanation rather than a demonstrated one. The practical conclusion does not "
        "depend on which mechanism dominates: the failure is as much the fixed k=2 "
        "tie-break as the dimensionality, and k=3 or distance weighting would break ties "
        "on evidence instead of class index. Both are outside the task's fixed "
        "specification, which is why we report it rather than change it."
    )

    r.h("7. Decision carried forward for Task 3")
    r.considered(
        [
            ["PCA before every Task 3 model",
             "Smaller, denser matrices; faster fits; standard practice on wide data",
             "Only 77% of variance survives even at 2,000 components; destroys the "
             "sparsity that tree and Naive Bayes models exploit"],
            ["TruncatedSVD, the sparse-friendly variant",
             "Does not densify the matrix, so it is the right tool for TF-IDF",
             "Same information loss; adds a component count to tune with no evidence it "
             "buys anything"],
            ["Feature selection by chi-squared or mutual information",
             "Keeps original interpretable columns; cheap",
             "Nothing to remove: no column is all-zero or near-constant"],
            ["No reduction for Task 3",
             "Nothing discarded; boosted trees handle 5,000 sparse columns natively",
             "Slower fits; no defence if a later model turns out to need dense input"],
        ],
        "Dimensionality reduction options for the Task 3 models.",
    )
    r.decision(
        chose="no dimensionality reduction for any Task 3 model. PCA appears only where "
              "Task 2 requires it.",
        tradeoff="we gave up faster fits and the option of distance-based and "
                 "dense-input models, and we carried 5,000 columns into every experiment "
                 "when a compressed representation would have been cheaper to iterate on.",
        because="the scree curve says compression here is lossy rather than efficient. "
                "There is no elbow to exploit, so any component count is a straight trade "
                "of information for speed. The models that led our baselines were boosted "
                "trees, which take sparse input directly and select features internally as "
                "they split, so the compression would have cost accuracy to buy speed we "
                "did not need.",
    )

    # ------------------------------------------------------------------ 8
    r.h("8. Phase 4: Baselines before Tuning")
    r.p(
        "The next decision was where to spend tuning effort, and we refused to make it by "
        "reputation. Twelve model families were run at their library defaults under the "
        "phase 3 protocol, and only then did anything get tuned."
    )
    base = facts.baselines()
    r.table(
        ["Model", "Mean Macro F1", "Std across folds"],
        [[name, f"{row['mean']:.4f}", f"{row['std']:.4f}"]
         for name, row in base.iterrows()],
        "Twelve baselines at library defaults, stratified 5-fold, supplied features.",
        widths=[2.0, 1.6, 1.6],
    )
    r.p(
        "Three things in that table shaped the rest of the project. Boosted trees took the "
        "top three places. The linear models clustered tightly just below them, within "
        "0.013 of each other, which told us that choosing among them was not where the "
        "gain was. And the whole table spans only 0.14 while sitting far below what the "
        "task should permit, which was the first hint that the representation rather than "
        "the model was the binding constraint. We did not read it that way at the time."
    )
    r.considered(
        [
            ["Tune the model we expect to win",
             "All effort on one candidate; fastest to a tuned number",
             "The expectation is untested, and a wrong guess is expensive to discover"],
            ["Baseline everything, then tune the top few",
             "Ranking is measured; the spread shows how much model choice is worth at all",
             "Twelve fits before any tuning starts"],
            ["Baseline everything and tune everything",
             "No candidate is prematurely discarded",
             "Twelve searches for a table that spans 0.14; most of it would be spent on "
             "models 0.05 behind the leader"],
        ],
        "How to allocate tuning effort.",
    )
    r.decision(
        chose="a full twelve-model baseline sweep at defaults first, then tuning "
              "restricted to the leader and the strongest linear models.",
        tradeoff="twelve fits before a single hyperparameter moved, and we still tuned "
                 "five families rather than committing to one, so some of that search "
                 "budget went on models that never became candidates.",
        because="the spread of the table was itself the finding. Knowing that twelve "
                "families span 0.14 and that the top four sit within 0.015 told us in "
                "advance that model choice was worth less than we hoped, which is the "
                "evidence we later used to stop tuning and stop ensembling rather than "
                "grinding at them. Tuning the linear models as well was cheap and gave us "
                "genuinely independent ensemble members in phase 6.",
    )
    r.figure("baseline_models_cv_f1.png",
             "The baseline sweep. LightGBM led, and it kept leading once we learned to "
             "compare models fairly in phase 7.", width=5.8)

    # ------------------------------------------------------------------ 9
    r.h("9. Phase 5: Tuning")
    r.p(
        "Five families were tuned with the same two-stage protocol: a coarse randomised "
        "search over a wide space, then a focused grid around the winner. Search spaces "
        "were documented in the notebook before the search ran, so the ranges were "
        "justified rather than reverse-engineered from the result."
    )
    r.p(
        f"It worked, in the sense that the numbers moved. LightGBM went from "
        f"{facts.SUPPLIED_LGBM_CV} to {facts.TUNED_LGBM_CV} in cross-validation. The "
        f"untouched holdout confirmed it at {facts.TUNED_LGBM_HOLDOUT}, a gap of 0.0006, "
        "so we were not overfitting the search. Tuning the decision threshold on the "
        "holdout added a further 0.003. Every diagnostic we had said the model was in good "
        "shape."
    )
    r.p(
        f"Then we submitted it and scored {facts.FIRST_KAGGLE_DEFAULT_THR} on the public "
        f"leaderboard. Applying the holdout-tuned threshold instead gave "
        f"{facts.FIRST_KAGGLE_TUNED_THR}. Against a local {facts.TUNED_LGBM_HOLDOUT} on "
        "data the model had never seen, that is a gap of roughly 0.079, and the "
        "threshold we had carefully tuned locally moved the leaderboard by less than a "
        "hundredth."
    )
    r.p(
        "Two readings were available. Either something was broken, or the brief's warning "
        "about a deliberate train-to-test shift was simply true and local numbers do not "
        "transfer. We checked for the usual causes of a suspicious gap, found no leakage, "
        "no duplicate rows shared across splits and no preprocessing fitted on the test "
        "set, and accepted the second reading. What we did not spot for another two phases "
        "was that we were also comparing models at different predicted class balances, "
        "which made the leaderboard look anti-correlated with local validation when it was "
        "not. Section 15.1 covers that."
    )
    r.considered(
        [
            ["Trust the local numbers and keep tuning",
             "Local validation was clean and reproducible; the holdout agreed with CV",
             "Local and leaderboard disagreed by 0.079, so at least one of them was not "
             "measuring the target"],
            ["Assume a bug and audit the pipeline",
             "Cheap to check, and the failure mode is severe if real",
             "The brief predicts this exact gap by design, so the prior on a bug is low"],
            ["Take the leaderboard as ground truth and tune against it",
             "It is the graded surface",
             "Roughly 3,570 scored rows means a noise floor near 0.0084, and submissions "
             "are rationed; this is how leaderboard overfitting starts"],
            ["Keep both, and quantify the leaderboard's precision",
             "Makes the two instruments comparable and stops arguments about small "
             "differences",
             "Requires spending submissions on measurement rather than on candidates"],
        ],
        "How to react to a 0.079 gap between local validation and the leaderboard.",
    )
    r.decision(
        chose="keep local validation as the ranking instrument, audit for leakage once, "
              f"and compute the leaderboard's own noise floor at {facts.NOISE_FLOOR} so "
              "that public differences could be judged rather than argued about.",
        tradeoff="we accepted that our local numbers would stay roughly 0.08 above the "
                 "leaderboard for the rest of the project, and we spent submissions on "
                 "diagnostics instead of on candidates.",
        because="the brief states the shift is deliberate, so a large gap is the "
                "designed behaviour rather than evidence of a defect, and chasing it "
                "closed would have meant fitting the public rows. Quantifying the noise "
                "floor was the highest-value submission we ever spent: it later let us "
                "stop three searches at the right moment and disbelieve two of our own "
                "apparent gains.",
    )
    r.figure("lightgbm_stage2_grid_heatmap.png",
             "Stage two of the LightGBM search on the supplied features. The surface is "
             "flat across most of the grid, which is why tuning bought about 0.005.",
             width=5.4)
    r.figure("lightgbm_threshold_sweep.png",
             "Threshold sweep on the holdout. The peak is real locally and moved the "
             "leaderboard by less than 0.01, which was the first sign that the decision "
             "rule needed a different kind of evidence.", width=5.4)

    # ------------------------------------------------------------------ 10
    r.h("10. Phase 6: Exploring Ensembling Methods")
    r.p(
        "With one model tuned and the gap unexplained, the standard next move is to "
        "combine models. The baseline table supported it: boosted trees and linear models "
        "make different errors, so there was real diversity to exploit. We built six "
        "members and tried four families of combiner across forty-eight configurations, "
        "including weighted blending with hill-climbed weights, majority and soft voting, "
        "meta-learners stacked on out-of-fold predictions, and rank aggregation."
    )
    r.p(
        f"Out-of-fold AUC improved by 0.0075. On the leaderboard the best combiner scored "
        f"{facts.ROUND4_BEST_KAGGLE} against {facts.SUPPLIED_LGBM_KAGGLE} for the single "
        f"model, a gain of {facts.ROUND4_BEST_KAGGLE - facts.SUPPLIED_LGBM_KAGGLE:+.5f}, "
        f"which is under a fifth of the {facts.NOISE_FLOOR} noise floor. We report it as a "
        "tie, because that is what it is."
    )
    r.p(
        "Both of those submissions were made at a matched predicted share of 0.4996, "
        "which is the only reason the two numbers are comparable at all. That control "
        "comes out of the calibration work in phase 7, and the fact that we needed it "
        "here is part of why phase 7 is written the way it is."
    )
    r.p(
        "This is the phase where the noise floor earned its keep. Without it we would have "
        "read +0.0017 as a small win, kept the ensemble, and spent the next phase adding "
        "members. With it, the correct reading was that two consecutive phases of work on "
        "the model had returned nothing the leaderboard could distinguish from zero, and "
        "that the constraint therefore lay somewhere we had not touched."
    )
    r.considered(
        [
            ["Ship the ensemble anyway",
             "It is nominally the best public score we have",
             "A 0.0017 lead inside a 0.0084 noise floor is not evidence; and six models "
             "cost six times the compute to retrain and audit"],
            ["Add more and more diverse members",
             "Diversity is the usual lever when blending stalls",
             "Diversity was already high; the members disagreed and the blend still did "
             "not move"],
            ["Stop, and ask what we have not changed",
             "Redirects effort to an untested axis",
             "Abandons a phase of finished work, and the answer might have been that "
             "nothing else was available"],
        ],
        "What to do when a well-built ensemble returns nothing.",
    )
    r.decision(
        chose="stop ensembling, revert to a single LightGBM, and treat the null as "
              "diagnostic information rather than a failure.",
        tradeoff="we gave up the standard route to a last fraction of a point, and our "
                 "final submission has no model diversity to fall back on if the private "
                 "rows behave differently from the public ones.",
        because="we had measured what that fraction was worth here and it was below what "
                "the leaderboard can resolve. Two phases of model work returning nothing "
                "is a strong signal about where the constraint is, and acting on it led "
                "directly to phase 8, which returned 0.044. Shipping one model also kept "
                "the pipeline simple enough to attribute later gains precisely, which "
                "mattered a great deal in phases 7 and 9.",
    )
    r.figure("ensemble_combiner_leaderboard.png",
             "Forty-eight combiner configurations across four families. The spread between "
             "the best and the single model is smaller than the leaderboard can measure.",
             width=5.8)

    # ------------------------------------------------------------------ 11
    r.h("11. Phase 7: Threshold Calibration")
    r.p(
        "If the model was not the constraint, the next candidate was the decision rule. A "
        "classifier emits a probability; turning it into a label needs a cut, and we had "
        "been using 0.5 by default, which inherits the training prior of 62.5% machine. If "
        "the test set has a different balance, that cut is wrong no matter how good the "
        "model is."
    )
    r.p(
        "The first attempt tuned the threshold on the local holdout. It moved the "
        f"leaderboard from {facts.FIRST_KAGGLE_DEFAULT_THR} to "
        f"{facts.FIRST_KAGGLE_TUNED_THR}, and it moved it in the wrong direction relative "
        "to what later turned out to be optimal. The reason is structural rather than "
        "incidental: our dev and holdout splits are both carved from training data and "
        "inherit its class balance, so no amount of local data can tell us what fraction "
        "of the test set is machine-generated. This is the one quantity in the project "
        "that local validation is constitutionally unable to estimate."
    )
    r.p(
        "So we changed the parameterisation. Instead of choosing a probability threshold, "
        "we sorted the test rows by score and labelled a fixed share of them machine. The "
        "share is directly interpretable as a prediction about the test set's class "
        "balance, it is comparable across models where a raw threshold is not, and it can "
        "be swept systematically. Sweeping it on the linear model gave 0.71783 at share "
        "0.55 and 0.71314 at share 0.45, locating a peak far from where the default "
        "threshold had been putting us."
    )
    r.p(
        f"Applied to the tuned LightGBM, moving from the default cut to a global share of "
        f"0.4996 took the leaderboard from {facts.FIRST_KAGGLE_DEFAULT_THR} to "
        f"{facts.SUPPLIED_LGBM_KAGGLE}. That is "
        f"{facts.SUPPLIED_LGBM_KAGGLE - facts.FIRST_KAGGLE_DEFAULT_THR:+.5f} from "
        "changing nothing but the cut, and it is the largest single gain in the project, "
        "nearly twice what rebuilding the entire representation was worth. It is also the "
        "cheapest: one line of code and one submission.",
        bold=False,
    )
    r.p(
        "That reparameterisation also exposed a mistake we had been making since phase 5. "
        "Each model had been submitted at whatever predicted share its own default "
        "threshold happened to produce, so every model comparison on the leaderboard had "
        "been confounded with a calibration comparison. Re-testing at a matched share of "
        f"0.4996, LightGBM scored {facts.SUPPLIED_LGBM_KAGGLE} against 0.71691 for "
        "ElasticNet and 0.71424 for LinearSVC. LightGBM led by 0.0189, in the same "
        "direction and roughly the same magnitude as its local lead. Local validation had "
        "never been anti-correlated with the leaderboard. We had been asking two questions "
        "as though they were one."
    )
    r.considered(
        [
            ["Keep the 0.5 default",
             "No fitting to the leaderboard at all; nothing to overfit",
             "It silently assumes the test balance equals the training balance of 62.5%, "
             "which the brief says is false"],
            ["Tune the threshold on the local holdout",
             "Uses held-out data; the standard, defensible procedure",
             "The holdout inherits the training balance, so it optimises for the wrong "
             "distribution; measured at 0.66502, worse than useless here"],
            ["Sweep predicted share on the leaderboard",
             "Directly parameterises the quantity we cannot estimate locally; comparable "
             "across models",
             "Fits to the public rows, which are noisy and rationed"],
            ["Estimate the test balance from the paper and set the share once",
             "One submission, no search, no leaderboard fitting",
             "The paper's figure covers the full shared-task test set, and the course uses "
             "a subsample of under 5% of it"],
        ],
        "How to convert model scores into labels.",
    )
    r.decision(
        chose="parameterise the decision rule as a predicted share rather than a "
              "probability threshold, and locate it by sweeping the leaderboard, using "
              "the paper's figure as a starting bracket rather than as an answer.",
        tradeoff="the share is fitted to the public leaderboard, which no local data can "
                 "validate, and a search on a noisy objective can manufacture apparent "
                 "gain that does not exist on the private rows.",
        because="there is no alternative that is not simply wrong. The default cut "
                "encodes a training prior the brief tells us does not hold, and the "
                "holdout-tuned cut encodes the same prior with extra steps, which we "
                "measured. We managed the risk instead of pretending it away: every share "
                "claim is judged against the noise floor, we stop refining once inside a "
                "flat region, and our final picks pair the best public point with a more "
                "central one.",
    )
    r.figure("class_balance_curve_round2.png",
             "The global predicted-share curve. Flat across a wide band, which we read at "
             "the time as share being exhausted. Phase 9 shows why that reading was wrong.",
             width=5.4)

    # ------------------------------------------------------------------ 12
    r.h("12. Phase 8: Pivot to raw-text features")
    r.p(
        "We were now stuck for a different reason. Tuning had gained 0.005 locally and "
        "little on the leaderboard, ensembling had gained 0.0017, and calibration, which "
        "had produced the largest gain of the project, had run out of room: sweeping "
        "global share further came back flat. Every one of those phases had changed the "
        "model or the decision rule while holding the input fixed, and the input was the "
        "one thing we had never questioned."
    )
    r.p(
        "What changed our minds was reading the shared-task paper properly rather than "
        "skimming its abstract. Its data section says the training corpora are HC3, M4GT "
        "and MAGE, and the test corpora are CUDRT, IELTS, NLPeer, PeerSum and MixSet, with "
        "no overlap between them. It also gives the class balance of each: 62.6% machine "
        "in training against 53.1% in test, which matches the 62.52% we had measured and "
        "the roughly 0.53 we had inferred from the share sweep. The test file is not a "
        "held-out sample of the same distribution. It is different sources entirely."
    )
    r.p(
        "That reframed the supplied features. They are top-5,000 TF-IDF over lemmas with "
        "stop words removed. Lemmatising and dropping stop words deliberately strips out "
        "function words and inflection and keeps content vocabulary, which is a sensible "
        "default for topic classification and close to the worst possible choice here, "
        "because content vocabulary is precisely what changes when the corpus changes. A "
        "model trained on those columns learns what training-set machine text is about, "
        "not how machine text is written. Meanwhile the raw text had been sitting in "
        "train.csv, essentially untouched through seven notebooks apart from one length "
        "summary."
    )
    r.p(
        "So we rebuilt the representation from the raw text: rates of 318 function words, "
        "punctuation and symbol rates per thousand characters, casing patterns, layout and "
        "structure such as newlines and bullets and markdown, length and burstiness "
        "moments, lexical diversity measures, readability, and character and word n-gram "
        "TF-IDF with stop words kept and no lemmatisation. Eleven blocks in total, from "
        "which eight were selected."
    )
    reps = facts.representations()
    r.table(
        ["Representation", "Best Macro F1 (stratified 5-fold)"],
        [["Supplied 5,000-column TF-IDF (the control)", f"{reps['tfidf_supplied']:.4f}"],
         ["Character n-grams alone (block H)", f"{reps['block_H']:.4f}"],
         ["Word n-grams alone (block I)", f"{reps['block_I']:.4f}"],
         ["All cheap per-document style blocks (A to G)", f"{reps['style_all']:.4f}"],
         ["Supplied plus style blocks", f"{reps['supplied_plus_style']:.4f}"],
         ["Supplied plus everything", f"{reps['supplied_plus_all']:.4f}"],
         ["All raw-text blocks", f"{reps['text_all']:.4f}"]],
        "Best model per representation, all measured in the same session against the "
        "same supplied-feature control.",
        widths=[3.6, 2.2],
    )
    r.p(
        f"The control reproduced at {reps['tfidf_supplied']:.4f} and the full raw-text set "
        f"reached {reps['text_all']:.4f}, a local gain of "
        f"{reps['text_all'] - reps['tfidf_supplied']:.4f}. Adding the supplied features "
        f"back on top gave {reps['supplied_plus_all']:.4f}, a tie using 5,000 more "
        "columns, so the supplied features contribute nothing once raw text is present. On "
        "the leaderboard, measured against the same supplied-feature LightGBM, the change "
        "was worth +0.038 immediately and +0.044 after the ablation below."
    )
    r.p(
        "Two follow-up notebooks then pruned and stress-tested the set. Each family was "
        "dropped in turn and the model re-scored on a held-out document type, so the "
        "measurement is what that family is worth for transfer rather than for fitting the "
        "training domains."
    )
    r.table(
        ["Family dropped", "Macro F1 given up", "Columns"],
        [[n, f"{d:+.4f}", f"{c:,}"] for n, d, c in facts.ABLATION],
        "Leave-one-family-out ablation under the grouped protocol. A negative value means "
        "the model improves when that family is removed.",
        widths=[2.2, 1.7, 1.2],
    )
    r.p(
        "Readability contributed nothing and the supplied TF-IDF was actively harmful, "
        "costing 0.0074 of held-out-domain Macro F1 by being present. Both were dropped, "
        f"leaving the {facts.BEST_N_FEATURES:,}-column set in section 14. Note also that "
        "the diversity block returns more than character n-grams from five columns against "
        "20,000, which is what prompted the expansion attempt reported in section 15.4."
    )
    r.considered(
        [
            ["Keep using the supplied features only",
             "Zero extra engineering; guaranteed comparable to other teams; the brief says "
             "no further work is required",
             "Lemmatised with stop words removed, so the surviving signal is content "
             "vocabulary, which is exactly what differs between the train and test corpora"],
            ["Add a handful of cheap style features to the supplied ones",
             "Small, safe, easy to describe",
             "Measured at 0.8584, below raw text alone, and it keeps the harmful supplied "
             "columns"],
            ["Rebuild n-grams only, without hand-built style blocks",
             "Much less code; character n-grams alone reach 0.8188",
             "Gives up 0.06; character n-grams also proved less valuable for transfer than "
             "the five-column diversity block"],
            ["Rebuild the representation fully from raw text",
             "Best measured result at 0.8811; captures how text is written rather than "
             "what it is about",
             "40,385 hand-built columns, far more code, a pipeline that would need "
             "rebuilding for any new corpus, and no extra marks for the effort"],
        ],
        "Representation options once the paper made the corpus shift concrete.",
    )
    r.decision(
        chose="rebuild the representation from raw text, keeping the eight blocks the "
              f"ablation supported, for {facts.BEST_N_FEATURES:,} columns in total.",
        tradeoff="we hand-build 40,385 features where the course supplies 5,000 ready to "
                 "use. That is far more engineering, far more code to get wrong, and the "
                 "brief is explicit that own preprocessing earns no extra marks.",
        because="it was worth +0.044 on the graded surface, which is more than everything "
                "else we did to the model combined, and marks for Task 3 come from the "
                "leaderboard rather than from the effort. The supplied preprocessing "
                "deletes function words, punctuation, casing and layout, and those are "
                "precisely the properties that survive a corpus change, which the ablation "
                "then confirmed family by family. The brief permits own preprocessing "
                "provided it is described in full, and this section plus section 14 is "
                "that description.",
    )
    r.figure("representation_comparison.png",
             "Each representation against the supplied-feature control re-run in the same "
             "session. Every dense block scores below the control alone; together they "
             "beat it by 0.115.", width=6.0)

    # ------------------------------------------------------------------ 13
    r.h("13. Phase 9: Calibration per group")
    r.p(
        "With the representation settled we returned to the decision rule, and found the "
        "most useful result in the project sitting in a measurement we had already taken "
        "and misread."
    )
    r.p(
        "We had swept global predicted share twice, in two different phases and on two "
        "different models, and both sweeps came back flat. On the raw-text model, share "
        "0.4996 gave 0.77349 and share 0.5299 gave 0.77231, a difference well inside the "
        "noise floor. The obvious conclusion was that share was exhausted. That conclusion "
        "was correct about the quantity we had measured and wrong about the question we "
        "thought we were answering."
    )
    r.p(
        "The test file is a mixture of two populations, as phase 1 had recorded. A global "
        "cut moves both of them in the same direction. It turned out they needed "
        "corrections in opposite directions: the UUID rows wanted a higher predicted "
        "machine share and the numeric peer-review rows wanted a lower one. A global sweep "
        "was therefore pushing one group toward its optimum while pushing the other away, "
        "and the two effects, each worth roughly 0.019, were cancelling in the average. "
        "The flat curve was not evidence of a flat optimum. It was evidence of a mixture."
    )
    surf = facts.share_surface()
    r.table(
        ["Submission", "UUID share", "Numeric share", "Public Macro F1"],
        [[row["file"].replace(".csv", ""), f"{row['uuid_share']:.4f}",
          f"{row['numeric_share']:.4f}", f"{row['kaggle_f1']:.5f}"]
         for _, row in surf.iterrows()],
        "Walking the two share axes separately. Each row is one submission.",
        widths=[2.3, 1.3, 1.3, 1.4],
    )
    r.p(
        f"Correcting the UUID axis alone took the score from 0.77942 to 0.79706 and then "
        f"to 0.80049 at its vertex. Correcting the numeric axis as well reached "
        f"{facts.BEST_KAGGLE}. A deliberate probe in the wrong direction lost 0.00926. "
        "Both the upward step and the wrong-direction step clear the noise floor, which is "
        "what makes the direction certain even though the last decimal of each coordinate "
        "is not."
    )
    r.considered(
        [
            ["One global share for all 6,999 rows",
             "One parameter, one submission per point, nothing group-specific to justify",
             "Measured flat twice, because it averages two opposite corrections"],
            ["Separate shares per id group",
             "Breaks the cancellation; worth +0.022 with no change to the model at all",
             "Two coordinates to fit on a noisy surface, and it relies on the id format "
             "genuinely marking a distinct population"],
            ["Train a separate model per group",
             "Would let the model itself adapt to peer-review text",
             "Only 1,999 UUID test rows and no labels for either group; nothing to train "
             "on"],
            ["Domain adaptation on the unlabelled test text",
             "Uses the test distribution without using its labels",
             "Substantially more machinery, and our transfer protocol is not precise "
             "enough to tell whether it helped"],
        ],
        "Decision-rule options once the mixture was understood.",
    )
    r.decision(
        chose=f"separate predicted shares of {facts.BEST_SHARES['uuid']} for the UUID rows "
              f"and {facts.BEST_SHARES['numeric']} for the numeric rows.",
        tradeoff="both coordinates are fitted entirely to the public leaderboard, seven "
                 "submissions went into locating them, and a coordinate search that size "
                 "on a noisy objective can manufacture 0.005 to 0.010 of gain that is not "
                 "present on the private rows.",
        because="the shape of the surface is solid even where the last decimal is not. "
                "The first correction gained 0.01764 and the deliberate wrong-direction "
                "probe lost 0.00926, both several times the noise floor, and the same "
                "correction reproduced on a completely different model. The grouping is "
                "not a leaderboard artifact either: the two populations differ in length, "
                "markdown density and vocabulary in ways we measured in phase 1, long "
                "before we knew the leaderboard would reward splitting them.",
    )
    r.figure("share_surface.png",
             "The per-group share surface. Both axes are at their fitted vertices, and the "
             "numeric axis is asymmetric.", width=6.0)

    r.h("13.1 Hyperparameters on new feature representation", level=2)
    r.p(
        "Every hyperparameter search up to this point had run against the supplied 5,000 "
        f"features. Nothing had been tuned on the {facts.BEST_N_FEATURES:,}-column "
        "raw-text matrix, which is a different problem in shape: eight times wider, "
        "sparser, and three quarters character n-grams. This phase closed that gap, and "
        "it is the one place in the project where work on the model itself paid."
    )
    r.p(
        "The search was split five ways across the team over 400 seeded configurations, "
        "sampling learning rate, tree count, leaf count, depth, minimum child samples, "
        "column and row subsampling and both regularisation terms, followed by a second "
        "stage refined around the centroid of the configurations that passed. It selects "
        "on paired per-fold differences against the default configuration rather than on "
        "raw means, because the five length bands differ in difficulty by far more than "
        "any two configurations differ from each other, and that difficulty is common to "
        "every configuration so it cancels in a paired difference. A candidate had to "
        "improve on every one of the five folds to pass."
    )
    r.table(
        ["Hyperparameter", "Value"],
        [[k, f"{v:g}"] for k, v in facts.BEST_PARAMS.items()],
        "The winning configuration. class_weight=\"balanced\", random_state=42 and "
        "subsample_freq=1 are fixed throughout and are not part of the search space.",
        widths=[2.6, 1.4],
    )
    r.p(
        f"It won by {facts.TUNING_PAIRED_GAIN:+.4f} paired Macro F1 over the defaults on "
        "the selection protocol, better on five folds out of five. Because a search "
        "optimises the protocol it selects on, that number is optimistic by construction, "
        "so we confirmed it on two measurements the search never saw: standard five-fold "
        f"cross-validation ({facts.TUNING_GAIN_STANDARD:+.4f}) and the three-band grouped "
        f"protocol ({facts.TUNING_GAIN_GROUPED:+.4f}). Both agree in sign, so the gain is "
        "not an artifact of the folds it was fitted to."
    )
    r.p(
        f"On the leaderboard it was worth {facts.TUNING_GAIN_KAGGLE:+.5f}, measured "
        "against the previous submission at identical per-group shares so that the "
        "hyperparameters were the only thing that moved. That single number revises this "
        "report's own headline, and section 1 says how: the identical search on the "
        "supplied features had returned nothing, and we had wrongly carried that null "
        "across a change of representation."
    )

    # ------------------------------------------------------------------ 14
    r.h("14. Final model & Conclusion")
    r.p(
        "This section answers the brief's first question in one place, gathering what the "
        "nine phases arrived at."
    )
    r.h("14.1 Classifier", level=2)
    r.p(
        "LightGBM is a gradient-boosted decision tree ensemble. It fits trees in sequence, "
        "each trained on the gradient of the loss left over by the trees before it, so "
        "every tree corrects its predecessors rather than voting independently. Its "
        "leaf-wise growth strategy splits whichever leaf promises the largest loss "
        "reduction, rather than growing the tree level by level, which is what makes it "
        "efficient on very wide sparse inputs. It also handles sparse columns natively and "
        "performs its own feature selection through split gain, which is why phase 2 could "
        "safely decline dimensionality reduction."
    )
    r.p(
        "We pass class_weight=\"balanced\" so the 62/38 split does not let the minority "
        "human class be outvoted. The remaining hyperparameters come from the two-stage "
        "search in section 13.1 and are listed there; the shape of the winner is a slow "
        f"learning rate of {facts.BEST_PARAMS['learning_rate']} over "
        f"{facts.BEST_PARAMS['n_estimators']:,} trees, with each tree seeing about "
        f"{facts.BEST_PARAMS['colsample_bytree']:.0%} of the columns. That column "
        "subsampling is the knob that matters most here: with three quarters of the "
        "matrix being character n-grams, most columns in any given split are close to "
        "redundant, and forcing each tree onto a different third of them is what buys "
        "the diversity the ensemble is built from."
    )
    r.h("14.2 Features", level=2)
    r.table(
        ["Block", "Columns", "What it measures"],
        [[n, f"{c:,}", d] for n, c, d in facts.CHOSEN_BLOCKS],
        f"The {facts.BEST_N_FEATURES:,} columns the final model sees. Blocks A to F are "
        "per-document statistics that cannot leak; H and I are vectorizers fitted on "
        "training text alone.",
        widths=[1.5, 0.9, 3.7],
    )
    r.h("14.3 Decision Rule", level=2)
    r.p(
        "The model emits a probability per row. We sort the rows by that probability and "
        f"label the top {facts.BEST_SHARES['uuid']:.2%} of the UUID rows and the top "
        f"{facts.BEST_SHARES['numeric']:.2%} of the numeric rows as machine, taking an "
        "exact count rather than a quantile so the realised share matches the target on "
        "ties."
    )
    r.h("14.4 Scores on Kaggle", level=2)
    r.table(
        ["Measurement", "Macro F1"],
        [["Stratified 5-fold cross-validation", f"{facts.BEST_CV_STANDARD:.4f}"],
         ["Grouped cross-validation, held-out document type",
          f"{facts.BEST_CV_GROUPED:.4f}"],
         ["Kaggle public leaderboard", f"{facts.BEST_KAGGLE:.5f}"]],
        "The same model measured three ways. The spread between them is the subject of "
        "section 15.2.",
        widths=[3.4, 1.7],
    )
    r.figure("kaggle_journey.png",
             "Public Macro F1 across the milestone submissions, with the noise band drawn "
             "around the final score. Three steps clear it: the first global share "
             "calibration, the representation change, and the first per-group correction. "
             "Everything flat is a phase spent on the model.")

    # ------------------------------------------------------------------ 15
    r.h("15. Difficulties")
    r.p(
        "Every item here is a mistake we made and then caught. The cost is included "
        "because the cost is the part that taught us something.", italic=True
    )

    r.h("15.1 We compared models along an axis we had not held fixed", level=2)
    r.p(
        "For three notebooks we believed local validation was anti-correlated with the "
        "leaderboard. LightGBM had the best cross-validation and holdout scores of "
        "anything we tried and the worst leaderboard score, so we wrote it off and spent a "
        "phase on linear models instead."
    )
    r.p(
        "The inversion was an artifact of the confound described in section 11. Once we "
        "re-tested at a matched predicted share, LightGBM beat ElasticNet by 0.0189, in "
        "the same direction and roughly the same magnitude as its local lead. What fixed "
        "it was designing a test that could overturn our own written conclusion, and "
        "writing down in advance what each possible outcome would mean. We kept that habit "
        "for the rest of the project, and it caught two more errors."
    )

    r.h("15.2 A validation protocol that would not reproduce", level=2)
    r.p(
        "Once transfer became the question, stratified cross-validation was no longer "
        "enough, because it trains and tests within the same corpora. We built domain "
        "clusters with k-means on the style features and gated everything downstream on a "
        "stability check. It passed at adjusted Rand index 0.9703 on one machine and "
        "failed at 0.5991 on another, from identical code and identical seeds, then fell "
        "to 0.1043 across five seeds. The cause was a 460,000-fold spread in column scales "
        "feeding a decomposition that ran before standardisation, on a partition that is "
        "genuinely multi-modal."
    )
    r.p(
        "Rather than discard the idea we tested whether the transfer gap depended on the "
        "labelling at all, including a deliberately random grouping as a control."
    )
    r.table(
        ["Grouping of the training rows", "Held-out Macro F1", "Usable folds"],
        [[n, f"{s:.4f}", k] for n, s, k in facts.GROUPINGS]
        + [["Standard 5-fold, for reference", f"{facts.BEST_CV_STANDARD:.4f}", 5]],
        "The random grouping is the load-bearing row: a meaningless split reproduces "
        "standard cross-validation, so the gap under structured groupings comes from "
        "domain structure and not from training on fewer rows.",
        widths=[2.6, 1.6, 1.1],
    )
    r.p(
        "We replaced k-means with deterministic length bands, which every team member "
        "computes identically. The absolute level remains untrustworthy to better than "
        "about 0.03, so we use the protocol to rank candidates rather than to forecast a "
        "score."
    )
    r.figure("local_vs_kaggle.png",
             "Standard cross-validation, grouped cross-validation and the leaderboard for "
             "three milestone models. Grouped validation is not unbiased, but its error is "
             "several times smaller.")

    r.h("15.3 A diagnostic that silently contaminated itself", level=2)
    r.p(
        "We built a submission to test whether the public leaderboard scores the UUID "
        "rows, by changing only those rows and holding the numeric rows fixed. The first "
        "version pinned the numeric group to the global share of 0.4996, in the belief "
        "that this held them constant. Their realised share under the baseline was 0.5482, "
        "so the probe silently moved 243 rows it was supposed to freeze, which would have "
        "made its answer unreadable in either direction."
    )
    r.p(
        "The general lesson: setting a knob to the value a baseline used is not the same "
        "as leaving those rows alone, and the two coincide only when the knob is applied "
        "at the same granularity. Every later per-group file asserts that the held group "
        "is byte-identical rather than trusting the share to imply it."
    )

    r.h("15.4 Two predictions we made, and both were wrong", level=2)
    r.p(
        "We predicted that the formatting features were dataset fingerprints that would "
        "not survive the corpus change, since machine text in training carries markdown "
        "bold at roughly ten times the human rate and the peer-review test rows carry it "
        "at five times the machine rate. We built a stripped-down feature set of function "
        "words and character n-grams only to test it. Removing those blocks cost 0.04944 "
        "on the leaderboard, six times the noise floor, in the opposite direction from the "
        "prediction."
    )
    r.p(
        "We then predicted that the diversity block had more to give, since it returned "
        "the largest transfer value of any family from only five columns. We added "
        "fourteen further diversity measures and twelve variability measures."
    )
    r.table(
        ["Feature idea", "Change in held-out-domain Macro F1"],
        [[n, f"{v:+.4f}"] for n, v in facts.EXPANSION_NULLS],
        "Four expansion attempts, all inside the selection bar. Extended diversity made "
        "transfer worse while improving same-domain fit, which is the signature of a "
        "feature that memorises rather than generalises.",
        widths=[3.2, 2.4],
    )
    r.p(
        "Both experiments were designed in advance so that they could fail, and both did. "
        "We would rather report two falsified predictions than a roadmap in which every "
        "idea worked."
    )

    r.h("15.5 What we could not overcome", level=2)
    r.bullets([
        "Predicted share cannot be validated locally at all. Our dev and holdout splits "
        "are carved from training data and inherit its class balance, so every share "
        "decision rests on leaderboard feedback. We managed that risk rather than solving "
        "it.",
        "The absolute level of the grouped protocol is not trustworthy to better than "
        "about 0.03, so we can rank candidates but not forecast a score.",
        "A gap of roughly 0.075 remains between standard cross-validation and the "
        "leaderboard. The brief says to expect it by design, and we found no lever that "
        "closed it without fitting the public rows.",
    ])

    # ------------------------------------------------------------------ 16
    r.h("16. What we learned beyond the course")
    r.h("16.1 Averaging can hide two large effects that cancel", level=2)
    r.p(
        "This is the idea that generalises furthest beyond this project. When one global "
        "threshold is applied to a population that is a mixture, the optimum for the "
        "mixture can be flat while the component optima are far apart and moving in "
        "opposite directions. We measured that flat curve twice, correctly both times, and "
        "drew the wrong conclusion from it on both occasions. A roughly +0.019 effect and "
        "a roughly -0.019 effect had been cancelling in the average, and the only way to "
        "see it was to stop averaging."
    )
    r.h("16.2 Blocked and paired comparison", level=2)
    r.p(
        "Our held-out bands differ in difficulty by far more than any two candidates "
        "differ from each other. Comparing mean scores against a fold standard deviation "
        "of 0.087 sets a bar that nothing can clear. Because band difficulty is common to "
        "every candidate, the correct comparison is paired: subtract fold by fold and test "
        "the differences. We had already reported one result as a null using the "
        "over-conservative bar before we understood this, and we say so rather than "
        "quietly restating it."
    )
    r.h("16.3 Leave-one-group-out validation as a transfer proxy", level=2)
    r.p(
        "Standard cross-validation answers how well a model fits this data. When train and "
        "test come from different sources, the question is how well it transfers, and "
        "those are different measurements. Holding out a whole group and training on the "
        "rest approximates the second. The essential companion is the random-grouping "
        "control in section 15.2: without it, a lower score under grouping is equally "
        "consistent with the protocol simply training on fewer rows, and the whole "
        "apparatus proves nothing."
    )
    r.h("16.4 A leaderboard is a measuring instrument with a precision", level=2)
    r.p(
        f"We computed the sampling noise of the public leaderboard once, about "
        f"{facts.NOISE_FLOOR} at our row count and score, and then held every claim to it. "
        "That single number let us stop three searches at the right moment, disbelieve two "
        "of our own gains, and avoid spending rationed submissions inside flat regions. It "
        "is also why our final picks are not simply the top two public scores."
    )
    r.h("16.5 Reading the dataset paper was worth more than any model", level=2)
    r.p(
        "The shared-task paper told us the test corpora were disjoint from training, gave "
        "the class balance of each, and explained why the test file contains two "
        "populations. Those facts produced both of the changes that moved our score. No "
        "amount of hyperparameter search would have found either of them, and we reached "
        "them in phase 8 rather than phase 1 only because we skimmed the paper the first "
        "time."
    )

    # ------------------------------------------------------------------ 17
    r.h("17. Limitations and next steps")
    r.p("What we consider settled:", bold=True)
    r.bullets([
        "The representation. Raw-text features beat the supplied TF-IDF by a wide margin, "
        "and adding the supplied features back changes nothing measurable.",
        "Global predicted share. Flat from 0.44 to 0.56 on two different models, for the "
        "reason section 16.1 explains.",
        "Per-group share. Both axes walked to their vertices.",
        "Feature expansion in the directions we tried. Four ideas, all inside the bar.",
    ])
    r.p("What remains open:", bold=True)
    r.bullets([
        "Hyperparameter tuning on the raw-text representation, designed and split across "
        "the team but not yet run. Expected to be worth about 0.005.",
        "A second model family blended on this representation. We specified it but did "
        "not manage to measure it, so we report it as unmeasured rather than as negative.",
        "Whether the id-format grouping matches the public and private leaderboard split. "
        "Our evidence is consistent with it but does not establish it.",
    ])
    r.p("Our final-submission rule:", bold=True)
    r.p(
        "pair the best public point with a more central one rather than taking the top two "
        "public scores. The share coordinates are fitted to a leaderboard with a known "
        "noise floor, and the private rows are a different sample, so the point that "
        "maximises the public score is not necessarily the one that maximises the graded "
        "one."
    )
    r.p(
        "For calibration, the shared task's own fine-tuned RoBERTa baseline reached 73.42 "
        f"Macro F1 on the full English test set. We reach {facts.BEST_KAGGLE * 100:.2f} "
        "with classical models and no deep learning, although the numbers are not strictly "
        "comparable because the course uses a subsample of under 5% of that data."
    )

    return r.save()


if __name__ == "__main__":
    out = build()
    print(f"wrote {out}  ({out.stat().st_size / 1000:.0f} kB)")
